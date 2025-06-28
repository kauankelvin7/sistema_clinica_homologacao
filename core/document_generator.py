from docx import Document
import os
from datetime import datetime
import subprocess

# Define o caminho para o arquivo do modelo
MODEL_PATH = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'models', 'modelo homologação.docx')
# Define o caminho para a pasta onde os documentos gerados serão salvos
OUTPUT_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'data', 'generated_documents')

# Garante que a pasta de saída exista
os.makedirs(OUTPUT_DIR, exist_ok=True)

def generate_document(data):
    """
    Carrega o modelo .docx, substitui os placeholders pelos dados fornecidos
    e salva o novo documento.
    """
    try:
        document = Document(MODEL_PATH)

        # Mapeamento dos campos do modelo para os dados recebidos
        replacements = {
            "{nome_paciente}": data.get("nome_paciente", ""),
            "{cpf_paciente}": data.get("cpf_paciente", ""), # AGORA PEGARÁ O CPF JÁ FORMATADO
            "{data_atestado}": data.get("data_atestado", ""),
            "{qtd_dias_atestado}": str(data.get("qtd_dias_atestado", "")),
            "{código_cid}": data.get("codigo_cid", ""),
            "{cargo_paciente}": data.get("cargo_paciente", ""),
            "{empresa_paciente}": data.get("empresa_paciente", ""),
            "___/___/____": datetime.now().strftime("%d/%m/%Y"),
            
            # Formatação do médico com tipo de registro
            "{nome_medico}{crm__medico}-{uf_crm_medico}": 
                f"{data.get('nome_medico', '')} {data.get('tipo_registro_medico', '')} {data.get('crm__medico', '')}-{data.get('uf_crm_medico', '')}."
        }

        # Iterar sobre todos os parágrafos do documento
        for paragraph in document.paragraphs:
            for key, value in replacements.items():
                if key in paragraph.text:
                    paragraph.text = paragraph.text.replace(key, value)
        
        # Iterar sobre todas as tabelas no documento
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for key, value in replacements.items():
                            if key in paragraph.text:
                                paragraph.text = paragraph.text.replace(key, value)

        # Gerar nome do arquivo de saída
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_filename = f"Declaracao_{data.get('nome_paciente', 'Paciente').replace(' ', '_')}_{timestamp}.docx"
        output_path = os.path.join(OUTPUT_DIR, output_filename)

        document.save(output_path)

        # --- ABRIR O ARQUIVO AUTOMATICAMENTE ---
        try:
            if os.name == 'nt': # Para Windows
                os.startfile(output_path)
            elif os.uname().sysname == 'Darwin': # Para macOS
                subprocess.Popen(['open', output_path])
            else: # Para Linux
                subprocess.Popen(['xdg-open', output_path])
            print(f"Abrindo arquivo: {output_path}")
        except Exception as e:
            print(f"Não foi possível abrir o arquivo automaticamente: {e}")
            print("Por favor, abra-o manualmente em:", output_path)

        return output_path

    except Exception as e:
        print(f"Erro ao gerar documento: {e}")
        return None